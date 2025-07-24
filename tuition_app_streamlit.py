import streamlit as st
import pandas as pd
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import gspread
import json
import base64
import mimetypes
import hashlib
import plotly.express as px
import io

from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from googleapiclient.errors import HttpError

# === CONFIGURATION ===
st.set_page_config(layout="wide", page_title="PRK Home Tuition")
UPI_ID = "9685840429@pnb"
SUBSCRIPTION_DAYS = 30
DATE_FORMAT = "%Y-%m-%d"

# === GOOGLE IDs ===
HOMEWORK_FOLDER_ID = "1e83Kseh47VMiKep7DKdOHr9ciwrbMyiO"
NOTEBOOK_FOLDER_ID = "1e9UpIdbkAw6AnUa3xAixIcdih5K9TRAH"
RECEIPT_FOLDER_ID = "1e9se9uNbpjdFGhzzOi1GdhYRcG4ZOmxn"

# === AUTHENTICATION ===
scopes = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]
try:
    decoded_creds = base64.b64decode(st.secrets["google_service"]["base64_credentials"])
    credentials_dict = json.loads(decoded_creds)
    credentials = Credentials.from_service_account_info(credentials_dict, scopes=scopes)
    client = gspread.authorize(credentials)
    drive_service = build("drive", "v3", credentials=credentials)
except Exception as e:
    st.error("Error connecting to Google APIs. Please check credentials and sharing permissions.")
    st.stop()

# === GOOGLE SHEETS ===
try:
    STUDENT_SHEET = client.open_by_key("10rC5yXLzeCzxOLaSbNc3tmHLiTS4RmO1G_PSpxRpSno").sheet1
    TEACHER_SHEET = client.open_by_key("1BRyQ5-Hv5Qr8ZnDzkj1awoxLjbLh3ubsWzpXskFL4h8").sheet1
    HOMEWORK_QUESTIONS_SHEET = client.open_by_key("1fU_oJWR8GbOCX_0TRu2qiXIwQ19pYy__ezXPsRH61qI").sheet1
    MASTER_ANSWER_SHEET = client.open_by_key("16poJSlKbTiezSG119QapoCVcjmAOicsJlyaeFpCKGd8").sheet1
except Exception as e:
    st.error(f"Could not open Google Sheets. Ensure all Sheet Keys are correct and shared with the service account: {e}")
    st.stop()

# === UTILITY FUNCTIONS ===
def make_hashes(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def check_hashes(password, hashed_text):
    if hashed_text:
        return make_hashes(password) == hashed_text
    return False

def upload_to_drive(path, folder_id, filename):
    try:
        mime_type, _ = mimetypes.guess_type(path)
        if mime_type is None:
            mime_type = 'application/octet-stream'
        media = MediaFileUpload(path, mimetype=mime_type, resumable=True)
        metadata = {"name": filename, "parents": [folder_id]}
        file = drive_service.files().create(body=metadata, media_body=media, fields="id, webViewLink", supportsAllDrives=True).execute()
        return file.get('webViewLink')
    except HttpError as error:
        st.error(f"An error occurred while uploading to Google Drive: {error}")
        return None

def create_answer_docx(student_name, student_class, answers_df):
    doc = Document()
    doc.add_heading('PRK Home Tuition - Submitted Answers', 0)
    doc.add_paragraph(f"**Student:** {student_name}")
    doc.add_paragraph(f"**Class:** {student_class}")
    doc.add_paragraph(f"**Date Generated:** {datetime.now().strftime(DATE_FORMAT)}")
    doc.add_paragraph()
    for (date, subject), group in answers_df.groupby(['Date', 'Subject']):
        doc.add_heading(f"Subject: {subject} (Assignment Date: {date})", level=2)
        for i, row in group.iterrows():
            p = doc.add_paragraph()
            p.add_run('Question: ').bold = True
            p.add_run(row.get('Question', ''))
            p_ans = doc.add_paragraph()
            p_ans.add_run('Your Answer: ').italic = True
            p_ans.add_run(row.get('Answer', ''))
            doc.add_paragraph()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def load_data(sheet):
    """
    Loads all data from a Google Sheet and correctly assigns the first row as the header.
    This is more robust than get_all_records().
    """
    all_values = sheet.get_all_values()
    if not all_values:
        return pd.DataFrame()
    
    headers = all_values[0]
    data = all_values[1:]
    df = pd.DataFrame(data, columns=headers)
    return df

def save_students_data(df):
    df_str = df.fillna("").astype(str)
    STUDENT_SHEET.clear()
    STUDENT_SHEET.update([df_str.columns.values.tolist()] + df_str.values.tolist())

def save_teachers_data(df):
    df_str = df.fillna("").astype(str)
    TEACHER_SHEET.clear()
    TEACHER_SHEET.update([df_str.columns.values.tolist()] + df_str.values.tolist())

def get_image_as_base64(path):
    """Converts an image file to a Base64 string."""
    try:
        with open(path, "rb") as f:
            data = f.read()
        encoded = base64.b64encode(data).decode()
        mime_type, _ = mimetypes.guess_type(path)
        if mime_type is None:
            mime_type = "image/png" # Default if type cannot be guessed
        return f"data:{mime_type};base64,{encoded}"
    except FileNotFoundError:
        return None

# === SESSION STATE INITIALIZATION ===
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user_name = ""
    st.session_state.user_role = ""
    st.session_state.user_gmail = ""

# === SIDEBAR & HEADER ===
st.sidebar.title("Login / Register")
if st.session_state.logged_in:
    st.sidebar.success(f"Welcome, {st.session_state.user_name}")
    if st.sidebar.button("Logout"):
        st.session_state.clear()
        st.rerun()

# --- Responsive Logo Section ---
prk_logo_b64 = get_image_as_base64("PRK_logo.jpg")
excellent_logo_b64 = get_image_as_base64("Excellent_logo.jpg")

if prk_logo_b64 and excellent_logo_b64:
    st.markdown(
        """
        <style>
        .logo-container {
            display: flex;
            justify-content: center;
            align-items: center;
            gap: 20px;
        }
        .logo-img {
            max-width: 45%;
            height: auto;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    st.markdown(
        f"""
        <div class="logo-container">
            <img src="{prk_logo_b64}" class="logo-img">
            <img src="{excellent_logo_b64}" class="logo-img">
        </div>
        """,
        unsafe_allow_html=True
    )
else:
    st.error("One or both logo files are missing.")

st.markdown("---")
# ------------------------------------

# === LOGIN / REGISTRATION ROUTING ===
if not st.session_state.logged_in:
    role = st.sidebar.radio("Login As:", ["Student", "Teacher", "Register", "Admin", "Principal"])

    if role == "Register":
        st.header("✍️ Registration")
        registration_type = st.radio("Register as:", ["Student", "Teacher"])
        
        if registration_type == "Student":
            with st.form("student_registration_form", clear_on_submit=True):
                name = st.text_input("Full Name")
                gmail = st.text_input("Gmail ID").lower().strip()
                cls = st.selectbox("Class", [f"{i}th" for i in range(6,13)])
                pwd = st.text_input("Password", type="password")
                if st.form_submit_button("Register (After Payment)"):
                    if not all([name, gmail, cls, pwd]):
                        st.warning("Please fill in all details.")
                    else:
                        df = load_data(STUDENT_SHEET)
                        if gmail in df["Gmail ID"].values:
                            st.error("This Gmail is already registered.")
                        else:
                            hashed_password = make_hashes(pwd)
                            till_date = (datetime.today() + timedelta(days=SUBSCRIPTION_DAYS)).strftime(DATE_FORMAT)
                            new_row = {"Sr. No.": len(df) + 1, "Student Name": name, "Gmail ID": gmail, "Class": cls, "Password": hashed_password, "Subscription Date": "", "Subscribed Till": till_date, "Payment Confirmed": "No"}
                            df_new = pd.DataFrame([new_row])
                            df = pd.concat([df, df_new], ignore_index=True)
                            save_students_data(df)
                            st.success("Registration successful! Waiting for payment confirmation.")
                            st.balloons()
            st.subheader("Payment Details")
            st.code(f"UPI ID: {UPI_ID}", language="text")
        
        elif registration_type == "Teacher":
            with st.form("teacher_registration_form", clear_on_submit=True):
                name = st.text_input("Full Name")
                gmail = st.text_input("Gmail ID").lower().strip()
                pwd = st.text_input("Password", type="password")
                if st.form_submit_button("Register Teacher"):
                    if not all([name, gmail, pwd]):
                        st.warning("Please fill in all details.")
                    else:
                        df_teachers = load_data(TEACHER_SHEET)
                        if gmail in df_teachers["Gmail ID"].values:
                            st.error("This Gmail is already registered as a teacher.")
                        else:
                            hashed_password = make_hashes(pwd)
                            new_row = {"Sr. No.": len(df_teachers) + 1, "Teacher Name": name, "Gmail ID": gmail, "Password": hashed_password, "Confirmed": "No"}
                            df_new = pd.DataFrame([new_row])
                            df_teachers = pd.concat([df_teachers, df_new], ignore_index=True)
                            save_teachers_data(df_teachers)
                            st.success("Teacher registered! Please wait for admin confirmation.")
                            st.balloons()
    else:
        st.header(f"🔑 {role} Login")
        with st.form(f"{role}_login_form"):
            login_gmail = st.text_input("Gmail ID").lower().strip()
            login_pwd = st.text_input("Password", type="password")
            if st.form_submit_button("Login"):
                sheet_to_check = TEACHER_SHEET if role in ["Admin", "Principal", "Teacher"] else STUDENT_SHEET
                name_col = "Teacher Name" if role in ["Admin", "Principal", "Teacher"] else "Student Name"

                df_users = load_data(sheet_to_check)
                user_data = df_users[df_users["Gmail ID"] == login_gmail]

                if not user_data.empty:
                    user_row = user_data.iloc[0]
                    if check_hashes(login_pwd, user_row.get("Password")):
                        can_login = False
                        if role == "Student":
                            if user_row.get("Payment Confirmed") == "Yes" and datetime.today() <= pd.to_datetime(user_row.get("Subscribed Till")):
                                can_login = True
                            else:
                                st.error("Subscription expired or not confirmed.")
                        elif role == "Teacher":
                            if user_row.get("Confirmed") == "Yes":
                                can_login = True
                            else:
                                st.error("Registration is pending admin confirmation.")
                        elif role in ["Admin", "Principal"]:
                            can_login = True

                        if can_login:
                            st.session_state.logged_in = True
                            st.session_state.user_name = user_row.get(name_col)
                            st.session_state.user_role = role.lower()
                            st.session_state.user_gmail = login_gmail
                            st.rerun()
                    else:
                        st.error("Invalid Gmail ID or Password.")
                else:
                    st.error("Invalid Gmail ID or Password.")

# === LOGGED-IN USER PANELS ===
if st.session_state.logged_in:
    current_role = st.session_state.user_role

    if current_role == "admin":
        st.header("👑 Admin Panel")
        tab1, tab2 = st.tabs(["Student Management", "Teacher Management"])
        with tab1:
            st.subheader("Manage Student Registrations")
            df_students = load_data(STUDENT_SHEET)
            st.markdown("#### Pending Payment Confirmations")
            unconfirmed_students = df_students[df_students.get("Payment Confirmed") != "Yes"]
            if unconfirmed_students.empty:
                st.info("No pending student payments.")
            else:
                for i, row in unconfirmed_students.iterrows():
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.write(f"**Name:** {row.get('Student Name')} | **Gmail:** {row.get('Gmail ID')}")
                    with col2:
                        if st.button("✅ Confirm Payment", key=f"confirm_student_{row.get('Gmail ID')}", use_container_width=True):
                            df_students.loc[i, "Subscription Date"] = datetime.today().strftime(DATE_FORMAT)
                            df_students.loc[i, "Subscribed Till"] = (datetime.today() + timedelta(days=SUBSCRIPTION_DAYS)).strftime(DATE_FORMAT)
                            df_students.loc[i, "Payment Confirmed"] = "Yes"
                            save_students_data(df_students)
                            st.success(f"Payment confirmed for {row.get('Student Name')}.")
                            st.rerun()
            st.markdown("---")
            st.markdown("#### Confirmed Students")
            confirmed_students = df_students[df_students.get("Payment Confirmed") == "Yes"]
            st.dataframe(confirmed_students)
        with tab2:
            st.subheader("Manage Teacher Registrations")
            df_teachers = load_data(TEACHER_SHEET)
            st.markdown("#### Pending Teacher Confirmations")
            unconfirmed_teachers = df_teachers[df_teachers.get("Confirmed") != "Yes"]
            if unconfirmed_teachers.empty:
                st.info("No pending teacher confirmations.")
            else:
                for i, row in unconfirmed_teachers.iterrows():
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.write(f"**Name:** {row.get('Teacher Name')} | **Gmail:** {row.get('Gmail ID')}")
                    with col2:
                        if st.button("✅ Confirm Teacher", key=f"confirm_teacher_{row.get('Gmail ID')}", use_container_width=True):
                            df_teachers.loc[i, "Confirmed"] = "Yes"
                            save_teachers_data(df_teachers)
                            st.success(f"Teacher {row.get('Teacher Name')} confirmed.")
                            st.rerun()
            st.markdown("---")
            st.markdown("#### Confirmed Teachers")
            confirmed_teachers = df_teachers[df_teachers.get("Confirmed") == "Yes"]
            st.dataframe(confirmed_teachers)

    elif current_role == "teacher":
        st.header(f"🧑‍🏫 Teacher Dashboard: Welcome {st.session_state.user_name}")
        df_homework = load_data(HOMEWORK_QUESTIONS_SHEET)
        st.subheader("Today's Submitted Homework")
        today_str = datetime.today().strftime(DATE_FORMAT)
        todays_homework = df_homework[(df_homework.get('Uploaded By') == st.session_state.user_name) & (df_homework.get('Date') == today_str)]
        if todays_homework.empty:
            st.info("You have not created any homework assignments today.")
        else:
            summary = todays_homework.groupby(['Class', 'Subject']).size().reset_index(name='Question Count')
            for index, row in summary.iterrows():
                st.success(f"Class: **{row.get('Class')}** | Subject: **{row.get('Subject')}** | Questions: **{row.get('Question Count')}**")
        st.markdown("---")
        create_tab, grade_tab, report_tab = st.tabs(["Create Homework", "Grade Answers", "My Reports"])
        with create_tab:
            st.subheader("Create a New Homework Assignment")
            if 'context_set' not in st.session_state:
                st.session_state.context_set = False
            if not st.session_state.context_set:
                with st.form("context_form"):
                    subject = st.selectbox("Subject", ["Hindi", "English", "Math", "Science", "SST", "Computer", "GK", "Advance Classes"])
                    cls = st.selectbox("Class", [f"{i}th" for i in range(6, 13)])
                    date = st.date_input("Date", datetime.today())
                    if st.form_submit_button("Start Adding Questions →"):
                        st.session_state.context_set = True
                        st.session_state.homework_context = {"subject": subject, "class": cls, "date": date}
                        st.session_state.questions_list = []
                        st.rerun()
            if st.session_state.context_set:
                ctx = st.session_state.homework_context
                st.success(f"Creating homework for: **{ctx['class']} - {ctx['subject']}** (Date: {ctx['date'].strftime(DATE_FORMAT)})")
                with st.form("add_question_form", clear_on_submit=True):
                    question_text = st.text_area("Enter a question to add:", height=100)
                    if st.form_submit_button("Add Question") and question_text:
                        st.session_state.questions_list.append(question_text)
                if st.session_state.questions_list:
                    st.write("#### Current Questions in this Assignment:")
                    for i, q in enumerate(st.session_state.questions_list):
                        st.write(f"{i + 1}. {q}")
                    if st.button("Final Submit Homework"):
                        rows_to_add = [[ctx['class'], ctx['date'].strftime(DATE_FORMAT), st.session_state.user_name, ctx['subject'], q_text] for q_text in st.session_state.questions_list]
                        HOMEWORK_QUESTIONS_SHEET.append_rows(rows_to_add, value_input_option='USER_ENTERED')
                        st.success("Homework submitted successfully!")
                        st.balloons()
                        del st.session_state.context_set, st.session_state.homework_context, st.session_state.questions_list
                        st.rerun()
                if st.session_state.context_set and st.button("Create Another Homework (Reset)"):
                    del st.session_state.context_set, st.session_state.homework_context, st.session_state.questions_list
                    st.rerun()
        with grade_tab:
            st.subheader("Grade Student Answers")
            df_all_answers = load_data(MASTER_ANSWER_SHEET)
            df_all_answers['Marks'] = pd.to_numeric(df_all_answers['Marks'], errors='coerce')
            ungraded_answers = df_all_answers[df_all_answers['Marks'].isna()]
            if ungraded_answers.empty:
                st.success("🎉 All submitted answers have been graded!")
            else:
                students_to_grade_gmail = ungraded_answers['Student Gmail'].unique().tolist()
                df_students = load_data(STUDENT_SHEET)
                gradable_students = df_students[df_students['Gmail ID'].isin(students_to_grade_gmail)]
                selected_student_name = st.selectbox("Select Student with Pending Answers", gradable_students['Student Name'].tolist())
                if selected_student_name:
                    student_gmail = gradable_students[gradable_students['Student Name'] == selected_student_name].iloc[0]['Gmail ID']
                    student_answers_df = df_all_answers[df_all_answers['Student Gmail'] == student_gmail]
                    for i, row in student_answers_df.sort_values(by='Date', ascending=False).iterrows():
                        st.markdown(f"**Date:** {row.get('Date')} | **Subject:** {row.get('Subject')}")
                        st.write(f"**Question:** {row.get('Question')}")
                        st.info(f"**Answer:** {row.get('Answer')}")
                        marks_value = row.get('Marks')
                        remarks_value = row.get('Remarks')
                        if pd.notna(marks_value):
                            st.success(f"**Graded: {GRADE_MAP_REVERSE.get(marks_value, 'N/A')} ({marks_value})**")
                            if pd.notna(remarks_value) and remarks_value:
                                st.warning(f"**Your Remark:** {remarks_value}")
                        else:
                            with st.form(key=f"grade_form_{i}"):
                                grade = st.selectbox("Grade", list(GRADE_MAP.keys()), key=f"grade_{i}")
                                remarks = st.text_area("Remarks/Feedback (Optional)", key=f"remarks_{i}")
                                if st.form_submit_button("Save Grade & Remarks"):
                                    cell_row = i + 2
                                    marks_col = df_all_answers.columns.get_loc('Marks') + 1
                                    remarks_col = df_all_answers.columns.get_loc('Remarks') + 1
                                    MASTER_ANSWER_SHEET.update_cell(cell_row, marks_col, GRADE_MAP[grade])
                                    MASTER_ANSWER_SHEET.update_cell(cell_row, remarks_col, remarks)
                                    st.success(f"Grade and remarks saved!")
                                    st.rerun()
                        st.markdown("---")
        with report_tab:
            st.subheader("My Reports")
            st.markdown("#### Class-wise Top 3 Students")
            df_answers_report = load_data(MASTER_ANSWER_SHEET)
            df_students_report = load_data(STUDENT_SHEET)
            if not df_answers_report.empty:
                df_answers_report['Marks'] = pd.to_numeric(df_answers_report['Marks'], errors='coerce')
                df_merged = pd.merge(df_answers_report, df_students_report, left_on='Student Gmail', right_on='Gmail ID')
                leaderboard = df_merged.groupby(['Class', 'Student Name'])['Marks'].mean().reset_index()
                top_students = leaderboard.groupby('Class').apply(lambda x: x.nlargest(3, 'Marks')).reset_index(drop=True)
                st.dataframe(top_students)

    elif current_role == "student":
        st.header(f"🧑‍🎓 Student Dashboard: Welcome {st.session_state.user_name}")
        df_students = load_data(STUDENT_SHEET)
        user_info_row = df_students[df_students["Gmail ID"] == st.session_state.user_gmail]
        if not user_info_row.empty:
            user_info = user_info_row.iloc[0]
            student_class = user_info.get("Class")
            st.subheader(f"Your Class: {student_class}")
            st.markdown("---")
            df_homework = load_data(HOMEWORK_QUESTIONS_SHEET)
            df_all_answers = load_data(MASTER_ANSWER_SHEET)
            homework_for_class = df_homework[df_homework.get("Class") == student_class]
            student_answers = df_all_answers[df_all_answers.get('Student Gmail') == st.session_state.user_gmail].copy()
            
            pending_tab, revision_tab = st.tabs(["Pending Homework", "Revision Zone"])
            with pending_tab:
                # Logic to find pending questions
                pass
            with revision_tab:
                # Logic to show graded answers
                pass
            
            st.subheader("Class Leaderboard")
            # Leaderboard logic here

    elif current_role == "principal":
        st.header("🏛️ Principal Dashboard")
        # Principal dashboard logic here
